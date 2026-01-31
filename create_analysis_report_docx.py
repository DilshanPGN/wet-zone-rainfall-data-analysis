"""
Create a simple Word document explaining the rainfall data analysis:
missing data filling, homogeneity test, ITA, and Mann-Kendall.
Run: python create_analysis_report_docx.py
Output: docs/Wet_Zone_Rainfall_Analysis_Report.docx
Equation images are rendered with matplotlib and saved to docs/equations/.
"""
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path(__file__).resolve().parent / "docs"
EQUATIONS_DIR = OUTPUT_DIR / "equations"
OUTPUT_FILE = OUTPUT_DIR / "Wet_Zone_Rainfall_Analysis_Report.docx"

# Matplotlib mathtext strings for each equation (use r'$...$')
EQUATIONS = {
    "missing_single": r"$P_X = \frac{\bar{P}_X^{(m)}}{\bar{P}_R^{(m)}} \times P_R$",
    "missing_multi": r"$P_X = \frac{1}{N} \sum_{i=1}^{N} \left( \frac{\bar{P}_X^{(m)}}{\bar{P}_{R_i}^{(m)}} \times P_{R_i} \right)$",
    "ita_1to1": r"$y = x \quad \text{(1:1 line, no trend)}$",
    "mk_S": r"$S = \sum_{i<j} \mathrm{sign}(x_j - x_i)$",
    "mk_var": r"$\mathrm{Var}(S) = \frac{n(n-1)(2n+5)}{18}$",
    "mk_z": r"$z = \frac{S \mp 1}{\sqrt{\mathrm{Var}(S)}}$",
    "sens_slope_ij": r"$\mathrm{slope}_{ij} = \frac{y_j - y_i}{t_j - t_i} \quad (i < j)$",
    "sens_beta": r"$\beta = \mathrm{median}\{\mathrm{slope}_{ij} : i < j\}$",
}


def render_equation(mathtext: str, filepath: Path, dpi: int = 150) -> None:
    """Render a mathtext equation to a PNG file."""
    fig, ax = plt.subplots(figsize=(6, 0.6))
    ax.axis("off")
    ax.text(0.5, 0.5, mathtext, fontsize=14, ha="center", va="center", transform=ax.transAxes)
    fig.patch.set_facecolor("white")
    fig.savefig(filepath, dpi=dpi, bbox_inches="tight", pad_inches=0.15, facecolor="white")
    plt.close(fig)


def add_equation_image(doc: Document, image_path: Path, width_inches: float = 4.5) -> None:
    """Add a centered equation image to the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    EQUATIONS_DIR.mkdir(exist_ok=True)

    # Render all equation images
    for key, mathtext in EQUATIONS.items():
        path = EQUATIONS_DIR / f"{key}.png"
        render_equation(mathtext, path)
    print(f"Equation images saved to: {EQUATIONS_DIR}")

    doc = Document()

    # Title
    doc.add_heading("Wet Zone Rainfall Data Analysis — Step-by-Step Guide", 0)
    doc.add_paragraph(
        "This document explains in simple terms how the rainfall data was prepared and "
        "analysed: filling missing data, checking homogeneity, and testing for trends. "
        "It is written so that anyone can follow the steps that led to the final results."
    )
    doc.add_paragraph()

    # ========== 1. MISSING DATA ==========
    doc.add_heading("1. Filling Missing Data", level=1)
    doc.add_paragraph(
        "Rainfall records often have gaps (e.g. a station did not report on some days). "
        "Before doing any trend analysis, those gaps must be filled in a standard way. "
        "This project uses the Normal Ratio Method, which is recommended by the World "
        "Meteorological Organization (WMO) for estimating missing rainfall using nearby stations."
    )
    doc.add_heading("Method: Normal Ratio Method", level=2)
    doc.add_paragraph(
        "The idea is simple: we use other stations (reference stations) that did record "
        "rain on the same day. We adjust their value by the long-term ratio between the "
        "target station and each reference station for that month, so seasonal patterns "
        "(e.g. wet vs dry months) are kept."
    )
    doc.add_paragraph("Single reference station (one station used to estimate):")
    add_equation_image(doc, EQUATIONS_DIR / "missing_single.png")
    doc.add_paragraph(
        "Where: P_X = estimated rainfall at the target station (mm); "
        "P_R = observed rainfall at the reference station on the same day (mm); "
        "P̄_X(m) = long-term mean rainfall at the target station for month m; "
        "P̄_R(m) = long-term mean rainfall at the reference station for month m; "
        "m = month (1–12)."
    )
    doc.add_paragraph("When several reference stations have data on that day, we average the estimates:")
    add_equation_image(doc, EQUATIONS_DIR / "missing_multi.png")
    doc.add_paragraph(
        "Where: N = number of reference stations with valid data on that day; "
        "Ri = reference station i; P_Ri = rainfall at station Ri on that day. "
        "The sum is over all reference stations that have data."
    )
    doc.add_heading("Fallback: Linear Interpolation", level=2)
    doc.add_paragraph(
        "If no reference station has data on that day (all are missing), we use linear "
        "interpolation: the missing value is set to the average of the previous and next "
        "known values in the same station's time series."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "In the Excel output, every cell that was filled by these methods is highlighted "
        "in light yellow so you can see which values are estimated."
    )
    doc.add_paragraph()

    # ========== 2. HOMOGENEITY TEST ==========
    doc.add_heading("2. Homogeneity Test", level=1)
    doc.add_paragraph(
        "A homogeneity test checks whether the rainfall series has a sudden step change "
        "(a break) that is not due to climate but to things like moving the station, "
        "changing the instrument, or changing the observer. If such a break exists, "
        "trend analysis can be misleading. So we test each station's annual rainfall "
        "totals for a single break."
    )
    doc.add_heading("What is homogeneity?", level=2)
    doc.add_paragraph(
        "Homogeneous means the series has no significant non-climatic break. "
        "Non-homogeneous means a significant break was detected: the mean level before "
        "and after a certain year is different in a way that is unlikely to be due to chance."
    )
    doc.add_heading("How the test is performed", level=2)
    doc.add_paragraph(
        "Daily rainfall is first summed by year for each station. Then two standard tests "
        "are run on each station's annual series (using the pyhomogeneity package):"
    )
    doc.add_paragraph(
        "• Pettitt test — A non-parametric test that finds the most likely single change-point "
        "and tests whether the mean before and after that point is significantly different. "
        "It is based on ranks (Mann–Whitney type), so it does not assume normal distribution."
    )
    doc.add_paragraph(
        "• SNHT (Standard Normal Homogeneity Test) — A parametric test that assumes the "
        "annual totals are roughly normally distributed. It also looks for a single shift "
        "in the mean and returns a change-point year and a p-value."
    )
    doc.add_paragraph(
        "Significance level used: α = 0.05. If the p-value is ≤ 0.05, we say the series is "
        "non-homogeneous (a break is detected). If p > 0.05, we say the series is homogeneous."
    )
    doc.add_heading("What the results mean", level=2)
    doc.add_paragraph(
        "• Homogeneous — No significant break at α = 0.05. The series can be used for "
        "trend analysis as one segment."
    )
    doc.add_paragraph(
        "• Non-homogeneous — A significant break was found. You may then correct the series "
        "(e.g. adjust the level after the break), or analyse the period before and after "
        "the break separately, or exclude the faulty period. The CSV and console output "
        "give the most probable break year for each test."
    )
    doc.add_paragraph(
        "In this project, all six stations (Colombo, Galle, Nuwara Eliya, Rathnapura, "
        "Maliboda, Deniyaya) were found to be homogeneous: all p-values were above 0.05, "
        "so no significant break was detected. The results are saved in homogeneity_test_results.csv."
    )
    doc.add_paragraph()

    # ========== 3. INNOVATIVE TREND ANALYSIS (ITA) ==========
    doc.add_heading("3. Innovative Trend Analysis (ITA)", level=1)
    doc.add_paragraph(
        "ITA is a simple graphical way to see if the second half of the period has higher "
        "or lower values than the first half. It does not assume any particular distribution "
        "(e.g. normal) for the data."
    )
    doc.add_heading("How ITA is performed", level=2)
    doc.add_paragraph("Step 1: Take the time series (e.g. annual rainfall for one station).")
    doc.add_paragraph(
        "Step 2: Split it into two halves. If there are n values: "
        "first half = first n/2 values, second half = remaining values. "
        "(If n is odd, the second half gets the middle point.)"
    )
    doc.add_paragraph(
        "Step 3: Plot a scatter: x-axis = values in the first half (in time order), "
        "y-axis = values in the second half (in time order). So the first point is "
        "(first value of first half, first value of second half), the second point is "
        "(second value of first half, second value of second half), and so on."
    )
    doc.add_paragraph(
        "Step 4: Draw the 1:1 line (the line y = x). This line means 'no change': "
        "second half = first half."
    )
    doc.add_heading("Equations and interpretation", level=2)
    doc.add_paragraph("Reference line (no trend):")
    add_equation_image(doc, EQUATIONS_DIR / "ita_1to1.png")
    doc.add_paragraph(
        "• Points above the 1:1 line → second-half value > first-half value → increasing trend."
    )
    doc.add_paragraph(
        "• Points below the 1:1 line → second-half value < first-half value → decreasing trend."
    )
    doc.add_paragraph(
        "• Mean change (summary number): mean(second half − first half). "
        "Positive means on average the second half is higher; negative means lower."
    )
    doc.add_paragraph(
        "The percentage of points above and below the 1:1 line is also reported. "
        "No equation is needed for that: it is just count above / total × 100 and "
        "count below / total × 100."
    )
    doc.add_paragraph(
        "Sen's slope (used for the trend line on the time series chart) is the median of "
        "all pairwise slopes between time points (see Section 4 for the formula)."
    )
    doc.add_paragraph()

    # ========== 4. MANN–KENDALL AND SEN'S SLOPE ==========
    doc.add_heading("4. Mann–Kendall Trend Test, Sen's Slope, and p-Value", level=1)
    doc.add_paragraph(
        "The Mann–Kendall (MK) test answers: 'Is there a statistically significant "
        "monotonic trend (consistently going up or down) in the series?' It does not "
        "assume normality; it uses the ranks of the data. Sen's slope gives the size of "
        "the trend (e.g. mm per year). The p-value tells us whether the trend is "
        "significant at the chosen level (here α = 0.05)."
    )
    doc.add_heading("Mann–Kendall: how it is performed", level=2)
    doc.add_paragraph(
        "For the series x₁, x₂, …, xₙ (in time order), we compare every pair (i, j) with i < j. "
        "We only look at the sign of the difference:"
    )
    doc.add_paragraph("Sign of (x_j − x_i):")
    doc.add_paragraph("  +1 if x_j > x_i")
    doc.add_paragraph("   0 if x_j = x_i")
    doc.add_paragraph("  −1 if x_j < x_i")
    doc.add_paragraph("The MK statistic S is the sum of these signs over all pairs (i, j) with i < j:")
    add_equation_image(doc, EQUATIONS_DIR / "mk_S.png")
    doc.add_paragraph(
        "• S > 0 → more increases than decreases → suggests increasing trend."
    )
    doc.add_paragraph(
        "• S < 0 → more decreases → suggests decreasing trend."
    )
    doc.add_paragraph(
        "• S = 0 → no trend."
    )
    doc.add_paragraph("Variance of S (when there are no ties) under 'no trend':")
    add_equation_image(doc, EQUATIONS_DIR / "mk_var.png")
    doc.add_paragraph("Where n = number of values. When there are ties, a correction is applied (as in the software).")
    doc.add_paragraph("Standardized statistic z (with continuity correction):")
    add_equation_image(doc, EQUATIONS_DIR / "mk_z.png")
    doc.add_paragraph("(Use S − 1 when S > 0, and S + 1 when S < 0.)")
    doc.add_paragraph(
        "Under 'no trend', z is approximately standard normal. The p-value is obtained "
        "from the standard normal distribution (two-tailed: we test for any trend, not "
        "only increase or only decrease)."
    )
    doc.add_heading("p-Value and trend decision", level=2)
    doc.add_paragraph(
        "We use significance level α = 0.05:"
    )
    doc.add_paragraph("• If p < 0.05 → reject 'no trend' → significant trend (increasing or decreasing).")
    doc.add_paragraph("• If p ≥ 0.05 → do not reject → no significant trend (reported as 'no trend').")
    doc.add_paragraph(
        "So 'no trend' in the output means 'no statistically significant monotonic trend "
        "at α = 0.05'. The slope can still be positive or negative; it is just not "
        "significant given the variability in the data."
    )
    doc.add_heading("Sen's slope", level=2)
    doc.add_paragraph(
        "Sen's slope is the median of all slopes between pairs of time points. For values "
        "y₁, y₂, …, yₙ at times t₁, t₂, …, tₙ (or indices 1, 2, …, n):"
    )
    doc.add_paragraph("Slope between pair (i, j), i < j:")
    add_equation_image(doc, EQUATIONS_DIR / "sens_slope_ij.png")
    doc.add_paragraph("Sen's slope β is the median of all such slopes:")
    add_equation_image(doc, EQUATIONS_DIR / "sens_beta.png")
    doc.add_paragraph(
        "For annual rainfall, β has units mm per year. It is robust to outliers. "
        "The trend line drawn in the app is: y = β × t + intercept, where the intercept "
        "is chosen so that the line fits the data in a robust way (e.g. median of (y_i − β t_i))."
    )
    doc.add_heading("Results in this project", level=2)
    doc.add_paragraph(
        "The Mann–Kendall results are in outputs/mann_kendall_results.csv. For all six "
        "stations, the test reported 'no trend': all p-values were above 0.05 (roughly "
        "0.25–0.87). Sen's slope is non-zero (positive for some stations, negative for "
        "others), but none of these slopes are statistically significant. So we conclude: "
        "there is no statistically significant monotonic trend in annual rainfall over "
        "the period covered by the data for any of the stations. ITA and MK together "
        "give a visual (ITA) and statistical (MK) picture of the same data."
    )
    doc.add_paragraph()

    # Summary — all steps in simple order
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Below are all the steps we followed to get the final results. They are done in this order."
    )
    doc.add_paragraph()

    def step(title: str, text: str) -> None:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        doc.add_paragraph(text)

    step("Step 1 — Fill missing data", (
        "Raw data had missing days. We filled them using nearby stations (Normal Ratio Method) "
        "or, when no station had data, by linear interpolation. The result was saved in an "
        "Excel file with filled cells highlighted in yellow."
    ))
    step("Step 2 — Check homogeneity", (
        "We checked if any station had a sudden break (e.g. instrument or location change). "
        "We used annual rainfall and two tests (Pettitt and SNHT). If a station was homogeneous, "
        "we kept it as one series. Results: homogeneity_test_results.csv."
    ))
    step("Step 3 — Innovative Trend Analysis (ITA)", (
        "We split each station's data into first half and second half of the period. We plotted "
        "second half vs first half. Above the 1:1 line = increase; below = decrease. We reported "
        "percentage above/below and mean change. This gives a simple picture of trend."
    ))
    step("Step 4 — Mann–Kendall and Sen's slope", (
        "We tested if there is a significant trend (up or down) in annual rainfall. We use p < 0.05: "
        "if p is below 0.05, the trend is significant; otherwise we say 'no trend'. We also computed "
        "Sen's slope (mm per year). Results: outputs/mann_kendall_results.csv."
    ))
    step("Final results", (
        "After these four steps: we have (1) a complete rainfall dataset, (2) all stations homogeneous, "
        "(3) ITA charts for visual trend, and (4) Mann–Kendall and Sen's slope. The final conclusion: "
        "no station had a significant trend at the 5% level."
    ))

    try:
        doc.save(OUTPUT_FILE)
        print(f"Report saved to: {OUTPUT_FILE}")
    except PermissionError:
        alt_file = OUTPUT_DIR / "Wet_Zone_Rainfall_Analysis_Report_new.docx"
        doc.save(alt_file)
        print(f"Could not overwrite (file may be open). Saved instead to: {alt_file}")


if __name__ == "__main__":
    main()
